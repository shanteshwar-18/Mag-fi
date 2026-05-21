import docx

doc = docx.Document(r'e:\Mag-Fi\Application 2\MagFi_Navigator_App2_PromptBook.docx')

output_lines = []
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style_name = para.style.name if para.style else "Unknown"
        output_lines.append(f'[P{i}] [{style_name}] {para.text}')

# Also check tables
output_lines.append("\n\n=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    output_lines.append(f"\n-- Table {t_idx} --")
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        output_lines.append(" | ".join(cells))

full_text = "\n".join(output_lines)

with open(r'e:\Mag-Fi\Application 2\prompt_book_extracted.txt', 'w', encoding='utf-8') as f:
    f.write(full_text)

print(f"Extracted {len(output_lines)} lines to prompt_book_extracted.txt")
print(f"Total characters: {len(full_text)}")
